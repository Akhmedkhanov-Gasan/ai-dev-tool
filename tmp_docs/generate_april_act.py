from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import Inches


SRC = Path(r"C:\Dev\ai-dev-tool\tmp_docs\Task_n_act_04_26_source.docx")
OUT = Path(r"C:\Dev\ai-dev-tool\tmp_docs\Task_n_act_05_26_april_2026.docx")


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text):
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    set_paragraph_text(paragraphs[0], text)


doc = Document(SRC)

paragraph_updates = {
    4: "Задание Заказчика №35 от 01.04.2026",
    9: "Общее количество часов (планово): 200",
    10: "Дата начала оказания Услуг: «01» Апреля 2026 г.",
    11: "Дата окончания оказания Услуг: «30» Апреля 2026 г.",
    13: "Срок представления результатов оказания услуг Заказчику:  \n"
        "Результаты оказания услуг необходимо предоставить не позднее 30.04.2026 ",
    14: "Стоимость Услуг (количество часов х стоимость часа труда): \n"
        "200 часов х 600 рублей = 120,000 рублей.",
    26: "Акт №35 от 30.04.2026 г.",
    29: "по Заданию №35 от 01.04.2026 Договора возмездного оказания услуг "
        "от 15.06.2023 № 15062023/01",
    30: "Общая стоимость услуг по заданию 120 000,00 рублей.",
    35: "Причитается к уплате Исполнителю по данному акту: 120 000 "
        "Сто двадцать тысяч рублей 00 копеек",
}

for index, text in paragraph_updates.items():
    set_paragraph_text(doc.paragraphs[index], text)

table = doc.tables[3]
table.rows[0].height = Inches(0.85)
table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
set_cell_text(table.cell(0, 2), "Период оказания работ/услуг,\n01.04.26-\n30.04.26")

set_cell_text(table.cell(2, 2), "90 ч")
set_cell_text(table.cell(2, 3), "\t54 000,00")

set_cell_text(table.cell(3, 2), "100 ч")
set_cell_text(table.cell(3, 3), "60 000,00")

set_cell_text(table.cell(6, 0), "Всего оказано услуг по настоящему акту                                                  200 ч")
set_cell_text(table.cell(6, 1), "Всего оказано услуг по настоящему акту                                                  200 ч")
set_cell_text(table.cell(6, 2), "Всего оказано услуг по настоящему акту                                                  200 ч")
set_cell_text(table.cell(6, 3), " 120 000,00")

doc.save(OUT)
print(OUT)
